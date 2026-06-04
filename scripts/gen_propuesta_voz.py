"""Genera propuesta_voz_callcenter.docx (one-shot)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)
NEGRO = RGBColor(0x00, 0x00, 0x00)
VERDE = RGBColor(0x2E, 0x7D, 0x32)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = AZUL
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = AZUL
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = GRIS
    return p


def parrafo(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.bold = bold
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        r = p.add_run(text)
    r.font.size = Pt(10.5)
    p.runs[0].text = text
    return p


def tabla(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False

    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], '1F4E79')

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9.5)
            if ri % 2 == 0:
                set_cell_bg(cell, 'F2F2F2')

    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return t


def link(doc, label, url):
    p = doc.add_paragraph()
    r1 = p.add_run(f'• {label}: ')
    r1.font.size = Pt(10)
    r1.bold = True
    r2 = p.add_run(url)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x06, 0x5F, 0xD2)
    r2.underline = True


def separador(doc):
    p = doc.add_paragraph()
    r = p.add_run('─' * 70)
    r.font.color.rgb = GRIS
    r.font.size = Pt(8)


def main():
    doc = Document()

    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    h1(doc, 'Propuesta Comercial')
    h2(doc, 'Sistema de Voz IA — Callcenter Inbound')

    info = doc.add_paragraph()
    info.add_run('Cliente: ').bold = True
    info.add_run('[Nombre del cliente]\n')
    info.add_run('Proveedor: ').bold = True
    info.add_run('[Tu nombre / empresa]\n')
    info.add_run('Fecha: ').bold = True
    info.add_run('2026-05-19\n')
    info.add_run('Inversión desarrollo: ').bold = True
    r = info.add_run('USD 5.000')
    r.bold = True
    r.font.color.rgb = VERDE
    info.add_run(' + mensualidad operativa')

    separador(doc)

    h2(doc, '1. Resumen ejecutivo')
    parrafo(doc,
        'Sistema de recepción telefónica con Inteligencia Artificial. El cliente '
        'final llama a un número de Ecuador (+593), un agente IA contesta en '
        'español, recolecta datos (cédula, nombre, intención), responde consultas '
        'sobre planes y servicios desde la base de conocimiento del cliente, y '
        'deriva a un asesor humano cuando corresponde. Integrado al CRM existente '
        '(Django + WhatsApp).')
    parrafo(doc, 'Capacidad estimada incluida: 10 llamadas/día × 10-15 min ≈ 3.000 min/mes.', bold=True)

    h2(doc, '2. Alcance del desarrollo (USD 5.000)')
    tabla(doc,
        ['Entregable', 'Descripción'],
        [
            ['Setup carrier + número EC', 'Compra DID Ecuador, configuración trunk SIP/TeXML, webhook'],
            ['Motor IVR conversacional', 'Modelos FlujoVoz/PasoVoz, admin Django, engine paso-a-paso'],
            ['Integración Media Streams', 'Refactor VozTwilioConsumer existente para producción'],
            ['Captura DTMF + voz', 'Híbrido teclado + reconocimiento natural (cédula, monto, opciones)'],
            ['Conexión con AgentesIA', 'Respuestas planes/FAQs con vectorstore existente'],
            ['Transferencia a asesor humano', 'Forward a celular o cola con softphone'],
            ['Grabación + transcripción', 'Almacena en S3, dashboard de llamadas'],
            ['1 flujo IVR completo', 'Saludo → menú → captura → agente IA → transfer'],
            ['Testing con llamadas reales', 'QA round-trip latencia, audio, edge cases'],
            ['Capacitación equipo cliente', '2 sesiones admin + manual operativo'],
            ['Soporte post-lanzamiento', '30 días incluidos'],
        ],
        col_widths=[5.5, 11.5])

    parrafo(doc, '')
    parrafo(doc, 'Cronograma: 6-8 semanas', bold=True)
    parrafo(doc, 'Pagos: 40% anticipo / 30% demo funcional / 30% entrega', bold=True)

    h2(doc, '3. Tecnologías que usaremos')

    h3(doc, 'Capa telefónica (carrier)')
    parrafo(doc,
        'Telnyx — proveedor SIP que vende números (DID) de Ecuador y soporta '
        'Media Streams (transmisión de audio en tiempo real vía WebSocket). '
        'Alternativas equivalentes: Twilio, Plivo, SignalWire.')

    h3(doc, 'STT — Speech to Text (captura de voz)')
    parrafo(doc,
        'Convierte el audio del cliente a texto en tiempo real. Opciones: '
        'Whisper local (gratuito, corre en tu servidor) o Deepgram cloud '
        '(streaming de baja latencia).')

    h3(doc, 'LLM — Modelo de lenguaje (cerebro)')
    parrafo(doc,
        'Entiende la intención del cliente y genera la respuesta. Opciones: '
        'Google Gemini Flash o OpenAI GPT-4o. Conectado al sistema AgentesIA '
        'existente con vectorstore (PDFs, FAQs, base de conocimiento del cliente).')

    h3(doc, 'TTS — Text to Speech (voz IA)')
    parrafo(doc,
        'Convierte la respuesta generada a audio natural. Opciones: Piper '
        'local (gratuito, voz robótica clara) o ElevenLabs cloud (voz natural '
        'indistinguible de humano).')

    h3(doc, 'Servidor y orquestación')
    parrafo(doc,
        'Django + Channels + Daphne (infraestructura existente del proyecto fastchat). '
        'Nginx con TLS para WebSocket seguro (wss://). Redis para colas de '
        'handoff y caché de sesión.')

    h3(doc, 'Almacenamiento')
    parrafo(doc,
        'PostgreSQL (existente): llamadas, transcripciones, métricas. '
        'S3 compatible (AWS / Wasabi / Cloudflare R2): grabaciones de audio.')

    h2(doc, '4. Comparativa de tiers tecnológicos')

    tabla(doc,
        ['Característica', 'BARATA', 'MEDIA (recomendada)', 'ALTA'],
        [
            ['Carrier', 'Telnyx EC', 'Telnyx EC', 'Telnyx EC'],
            ['STT', 'Whisper local', 'Deepgram Nova-2', 'Deepgram Nova-2'],
            ['LLM', 'Gemini Flash', 'Gemini Flash / GPT-4o-mini', 'GPT-4o'],
            ['TTS', 'Piper local', 'ElevenLabs Flash v2.5', 'ElevenLabs Pro / Cartesia'],
            ['Latencia turno', '2.5 - 4 seg', '0.6 - 1.2 seg', '0.4 - 0.9 seg'],
            ['Calidad voz', 'Robótica clara', 'Natural', 'Indistinguible humano'],
            ['Precisión STT español EC', '85 - 90%', '94 - 97%', '94 - 97%'],
            ['Barge-in (interrumpir IA)', 'No', 'Sí', 'Sí'],
            ['Llamadas concurrentes', '2 - 4', '20 - 40', '20 - 40'],
            ['Datos salen del país', 'No (local)', 'Sí (US)', 'Sí (US)'],
            ['Costo proveedor/mes (3k min)', 'USD 70', 'USD 110', 'USD 185'],
            ['Precio sugerido cliente', 'USD 300 - 350', 'USD 400 - 500', 'USD 650 - 850'],
            ['Margen neto/mes', 'USD 230 - 280', 'USD 290 - 390', 'USD 465 - 665'],
        ],
        col_widths=[5.0, 4.0, 4.0, 4.0])

    h2(doc, '5. Qué te permite cada tier')

    h3(doc, 'BARATA — Recepcionista básica')
    bullet(doc, 'Casos: citas simples, recordatorios, mensajes informativos.')
    bullet(doc, 'Ventaja: datos médicos/legales no salen del país (cumple normativa local).')
    bullet(doc, 'Limitación: cliente percibe pausas, voz claramente robótica, no se puede interrumpir.')
    bullet(doc, 'Tope práctico: ~3 clientes simultáneos sin escalar servidor.')

    h3(doc, 'MEDIA — Asistente conversacional profesional (RECOMENDADA)')
    bullet(doc, 'Conversación fluida, cliente puede interrumpir a la IA en cualquier momento.')
    bullet(doc, 'Vende activamente, califica leads, responde preguntas complejas sobre planes y servicios.')
    bullet(doc, 'Escala a 10-30 clientes finales sin tocar arquitectura.')
    bullet(doc, 'Único riesgo: dependencia de servicios cloud (mitigable con fallback Piper local).')

    h3(doc, 'ALTA — Indistinguible de humano')
    bullet(doc, 'Voz clonada del cliente, razonamiento avanzado GPT-4o.')
    bullet(doc, 'Nicho: bancos, clínicas premium, cuentas USA/EU.')
    bullet(doc, 'ROI sólo si el cliente final paga USD 800+/mes.')

    h2(doc, '6. Costo operativo mensual (proveedor)')
    parrafo(doc, 'Escenario: 1 cliente, 10 llamadas/día × 12.5 min = ~3.000 min/mes', bold=True)
    parrafo(doc, 'Tier MEDIA (recomendado):', bold=True)

    tabla(doc,
        ['Servicio', 'Tarifa', 'Mes'],
        [
            ['Telnyx DID Ecuador', 'fijo', 'USD 5'],
            ['Telnyx minutos entrantes EC', 'USD 0.020/min', 'USD 60'],
            ['Deepgram Nova-2 streaming', 'USD 0.0043/min', 'USD 13'],
            ['ElevenLabs Flash v2.5 (TTS)', '~USD 0.015/min IA', 'USD 27'],
            ['Google Gemini Flash (LLM)', '~USD 0.001/turno', 'USD 3'],
            ['AWS S3 grabaciones (~9 GB)', 'USD 0.023/GB', 'USD 1'],
            ['Servidor Daphne', '(ya existente)', 'USD 0'],
            ['TOTAL COSTO PROVEEDOR', '', 'USD ~109'],
        ],
        col_widths=[7.5, 5.0, 4.5])

    parrafo(doc, '')
    parrafo(doc, 'Cobro al cliente: USD 450/mes (plan estándar)', bold=True)
    p = doc.add_paragraph()
    r = p.add_run('Margen neto: ~USD 340/mes recurrente por cliente')
    r.bold = True
    r.font.color.rgb = VERDE

    h2(doc, '7. Mensualidad recurrente al cliente')

    tabla(doc,
        ['Plan', 'Mensualidad', 'Incluye'],
        [
            ['Recepcionista (tier media)', 'USD 450/mes', '1 número, 3.500 min, 1 flujo, soporte 48h'],
            ['Add-on flujo extra', '+USD 80/mes', 'Flujo IVR adicional'],
            ['Add-on voz premium', '+USD 120/mes', 'Upgrade ElevenLabs Pro'],
            ['Add-on integración CRM externo', '+USD 150/mes', 'API a sistema cliente'],
            ['Exceso minutos', 'USD 0.18/min', 'Sobre límite incluido'],
        ],
        col_widths=[5.5, 4.0, 7.5])

    parrafo(doc, '')
    bullet(doc, 'Contrato mínimo: 6 meses (para amortizar setup).')
    bullet(doc, 'Aumento anual: 8% indexado a inflación + mejoras servicio.')

    h2(doc, '8. Links oficiales — proveedores')

    h3(doc, 'Carriers (DID + minutos)')
    link(doc, 'Telnyx pricing', 'https://telnyx.com/pricing/call-control')
    link(doc, 'Telnyx cobertura global', 'https://telnyx.com/global-coverage')
    link(doc, 'Twilio Voice pricing', 'https://www.twilio.com/voice/pricing')
    link(doc, 'Plivo pricing', 'https://www.plivo.com/pricing/')
    link(doc, 'SignalWire pricing', 'https://signalwire.com/pricing')

    h3(doc, 'STT — Speech to Text')
    link(doc, 'Deepgram pricing', 'https://deepgram.com/pricing')
    link(doc, 'Deepgram docs', 'https://developers.deepgram.com')
    link(doc, 'AssemblyAI Realtime', 'https://www.assemblyai.com/pricing')
    link(doc, 'Whisper local (faster-whisper)', 'https://github.com/SYSTRAN/faster-whisper')

    h3(doc, 'LLM — Cerebro')
    link(doc, 'Google Gemini pricing', 'https://ai.google.dev/pricing')
    link(doc, 'Google AI Studio (API keys)', 'https://aistudio.google.com/apikey')
    link(doc, 'OpenAI API pricing', 'https://openai.com/api/pricing')
    link(doc, 'OpenAI platform', 'https://platform.openai.com')

    h3(doc, 'TTS — Voz')
    link(doc, 'ElevenLabs pricing', 'https://elevenlabs.io/pricing')
    link(doc, 'ElevenLabs API docs', 'https://elevenlabs.io/docs/api-reference')
    link(doc, 'Cartesia Sonic pricing', 'https://cartesia.ai/pricing')
    link(doc, 'Piper TTS local (GitHub)', 'https://github.com/rhasspy/piper')
    link(doc, 'Modelos Piper español', 'https://github.com/rhasspy/piper/blob/master/VOICES.md')

    h3(doc, 'Frameworks orquestación voz')
    link(doc, 'Pipecat (recomendado)', 'https://github.com/pipecat-ai/pipecat')
    link(doc, 'LiveKit Agents', 'https://docs.livekit.io/agents/')
    link(doc, 'Vocode', 'https://github.com/vocodedev/vocode-core')

    h3(doc, 'Storage / Infraestructura')
    link(doc, 'AWS S3 pricing', 'https://aws.amazon.com/s3/pricing/')
    link(doc, 'Cloudflare R2 (sin egress)', 'https://developers.cloudflare.com/r2/pricing/')
    link(doc, 'Wasabi pricing', 'https://wasabi.com/cloud-storage-pricing')

    h2(doc, '9. Garantías y soporte')
    bullet(doc, 'Uptime objetivo: 99% mensual (excluye caídas de carrier/cloud).')
    bullet(doc, 'Tiempo respuesta soporte: 24h hábiles email, 4h críticos.')
    bullet(doc, 'Backup automático: transcripciones + configuración en base de datos diaria.')
    bullet(doc, 'Soporte post-lanzamiento: 30 días incluidos en setup.')
    bullet(doc, 'Capacitación: 2 sesiones online + manual escrito.')

    h2(doc, '10. Fuera de alcance (cotización adicional)')
    bullet(doc, 'Llamadas salientes automáticas (marcador predictivo).')
    bullet(doc, 'Integración con sistemas legacy del cliente (ERP, otros CRMs).')
    bullet(doc, 'Flujos IVR adicionales más allá del primero incluido.')
    bullet(doc, 'Voces clonadas custom.')
    bullet(doc, 'Reportes BI / dashboards analíticos avanzados.')
    bullet(doc, 'Cumplimiento HIPAA / PCI (requiere infraestructura aparte).')

    separador(doc)

    h2(doc, '11. Aceptación')
    parrafo(doc, 'Firmando el presente documento, el cliente acepta el alcance, '
                  'cronograma y condiciones económicas detalladas.')
    parrafo(doc, '')
    parrafo(doc, '')
    parrafo(doc, '__________________________            __________________________')
    parrafo(doc, 'Cliente                                                                                Proveedor')
    parrafo(doc, 'Nombre:                                                                              Nombre:')
    parrafo(doc, 'CI / RUC:                                                                             CI / RUC:')
    parrafo(doc, 'Fecha:                                                                                 Fecha:')

    out = 'propuesta_voz_callcenter.docx'
    doc.save(out)
    print(f'OK -> {out}')


if __name__ == '__main__':
    main()
